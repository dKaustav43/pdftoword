import pdfplumber
from docx import Document
from pathlib import Path

def pdftoword(input_pdf_path:str, output_word_path:str, start:int, end:int):
    
    '''
    The following code takes a pdf as an input and converts to a docx file.
    Input the path to your pdf file into pdf_file. 
    Imput the path where you would like to store the converted text file. 
    '''
    input_path = Path(input_pdf_path)

    if not input_path.is_file():
        raise FileNotFoundError(f"Input file does not exist: {input_pdf_path}")
    
    with pdfplumber.open(input_path) as pdf:
        total_pages = len(pdf.pages) #counting the total number of pages

    if start<0 or end>=total_pages: #validating the start and end page request
        raise ValueError("Page range is out of bounds"
                         f"PDF has {total_pages} pages.")
    
    doc = Document()

    with pdfplumber.open(input_path) as pdf:
        for i, page in enumerate(pdf.pages[start-1:end]):
            text = page.extract_text()
            if text:
                doc.add_paragraph(f"--- Page {i+1} ---")
                doc.add_paragraph(text)

    doc.save(output_word_path)

    return str(output_word_path)

def main():

    input_pdf_path = input("Enter the path of the pdf file here: ")
    output_word_path = input("Enter the path where you would like to store the converted word file:")

    start = int(input("Start page index: "))
    end = int(input("End page index:"))

    result = pdftoword(input_pdf_path,output_word_path,start,end)

    print(f"Created{result}")

if __name__ == "__main__":
    main()
